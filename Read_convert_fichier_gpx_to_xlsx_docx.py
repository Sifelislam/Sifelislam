#-----------------------------------------read fichier GPX et convert to Excel & word---------------
#-----------------ariba sifelislam étudiants en école nationale supérieure des sciences géodésique et techniques spéciales-algérie-----------
import pandas as pd
from pandas import DataFrame 
import gpxpy
import docx
#import os

#Open Gpx file
gpx_file = open('C:\\Users\\MOHSIM\\Downloads\\fells_loop.gpx', 'r') 
gpx = gpxpy.parse(gpx_file) \

name=[]
latitude=[]
longitude=[]
elevation=[]

for waypoint in gpx.waypoints: 
	name.append(waypoint.name)
	latitude.append(waypoint.latitude)
	longitude.append(waypoint.longitude)
	elevation.append(waypoint.elevation)


#Create a dict 
d={
	"Name":name,
	"Latitude":latitude,
	"Longitude":longitude,
	"Elevation":elevation
}

#Create a DataFrame
pd=DataFrame(d,index=None)
df=pd.set_index("Name")

# print(df.index) Point Names 

for i in df.index: 
	(df.loc[i],"\n-------------------------------------\n")

choix=input('entrer xlsx ou doc')
if choix=='xlsx':
    #Convert to excel
    pd.to_excel("xxx.xlsx",index=None)
    print("xlsx")
elif choix=='doc':
    a=[]
    doc= docx.Document()
    doc.add_heading('Table Document',0)
    for i in range(0,len(latitude)):
        resultat=[i,name[i],str(latitude[i]),str(longitude[i]),str(elevation[i])]
        a.append(resultat)

    menuTable = doc.add_table(rows=1,cols=5)
    menuTable.style= 'Table Grid'
    hdr_cells = menuTable.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = 'NAME'
    hdr_cells[2].text = 'LATITUDE'
    hdr_cells[3].text = 'LONGITUDE'
    hdr_cells[4].text = 'ELEVATION'
    for ID , NAME ,LATITUDE,LONGITUDE,ELEVATION in a:
        row_cells = menuTable.add_row().cells
        row_cells[0].text= str(ID)
        row_cells[1].text= NAME
        row_cells[2].text= LATITUDE
        row_cells[3].text= LONGITUDE
        row_cells[4].text= ELEVATION
    doc.save('table.docx')
    #os.system('start table.docx')
    print('doc')
        
